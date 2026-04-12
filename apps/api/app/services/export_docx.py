from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import datetime

def generate_office_action_response_docx(
    draft_text: str,
    metadata: dict
) -> bytes:
    """
    Generate a professional USPTO Office Action Response DOCX.
    Follows standards: 1" margins, Times New Roman, Header Caption.
    """
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style defaults
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # 1. Header Caption Table
    table = doc.add_table(rows=1, cols=2)
    table.stylename = 'Table Grid'
    
    # Left side: Application info
    left_cell = table.cell(0, 0)
    left_p = left_cell.paragraphs[0]
    left_p.add_run(f"Applicant: {metadata.get('applicant', 'TBD')}\n")
    left_p.add_run(f"App. No.: {metadata.get('application_number', 'TBD')}\n")
    left_p.add_run(f"Filed: {metadata.get('filing_date', 'TBD')}\n")
    left_p.add_run(f"Title: {metadata.get('title', 'Untitled')}")

    # Right side: Examiner info
    right_cell = table.cell(0, 1)
    right_p = right_cell.paragraphs[0]
    right_p.add_run(f"Examiner: {metadata.get('examiner', 'TBD')}\n")
    right_p.add_run(f"Art Unit: {metadata.get('art_unit', 'TBD')}\n")
    right_p.add_run(f"Docket No.: {metadata.get('docket_number', 'TBD')}")

    doc.add_paragraph("\n")

    # 2. Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RESPONSE TO OFFICE ACTION")
    run.bold = True
    run.underline = True

    doc.add_paragraph("\n")

    # 3. Content
    # Split the draft_text by paragraphs or double newlines
    paragraphs = draft_text.split("\n")
    for p_text in paragraphs:
        if p_text.strip():
            p = doc.add_paragraph(p_text.strip())
            p.paragraph_format.space_after = Pt(12)

    # 4. Signature block
    doc.add_paragraph("\n\nRespectfully submitted,")
    doc.add_paragraph("\n__________________________")
    doc.add_paragraph("Registration No. __________")
    doc.add_paragraph(f"Date: {datetime.date.today().strftime('%B %d, %Y')}")

    # Save to buffer
    target_stream = io.BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()
