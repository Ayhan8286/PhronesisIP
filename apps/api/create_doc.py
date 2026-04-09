import docx

doc = docx.Document()
doc.add_heading('VoiceAI Hospital Assistant - Engineering Notes', 0)
doc.add_paragraph('Overview: We built a voice assistant for hospitals. It receives voice commands from nurses and doctors.')
doc.add_paragraph('Core features:\n1. A speech input module to receive medical voice commands.\n2. A pharmacological risk-weighting module.\n3. An intent routing engine.')
doc.add_paragraph('Why it is novel: Unlike Siri or Alexa, our system explicitly computes a pharmacological risk score when hearing drug names. If the risk score is high (e.g. hearing morphine vs aspirin), it requires an explicit user confirmation step before routing the intent to the hospital EMR system. This prevents accidental lethal dosing commands.')

doc.save(r'd:\box mation\apps\api\voiceai_spec.docx')
print("Document created.")
